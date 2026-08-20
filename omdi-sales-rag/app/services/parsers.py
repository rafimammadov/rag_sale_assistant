from __future__ import annotations

import csv
import io
import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from app.services.chunking import normalize_text


@dataclass(slots=True)
class ParsedSection:
    text: str
    page: int | None = None
    section: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


def _first_heading(text: str, fallback: str) -> str:
    lines = [normalize_text(line) for line in text.splitlines() if normalize_text(line)]
    return " - ".join(lines[:2])[:300] if lines else fallback


def parse_pdf(path: Path) -> list[ParsedSection]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    sections: list[ParsedSection] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = normalize_text(page.extract_text() or "")
        if text:
            sections.append(
                ParsedSection(
                    text=text,
                    page=page_number,
                    section=_first_heading(text, f"Page {page_number}"),
                )
            )
    if not sections:
        raise ValueError(
            "The PDF contains no extractable text. Run OCR first or upload a searchable PDF."
        )
    return sections


def parse_docx(path: Path) -> list[ParsedSection]:
    from docx import Document

    document = Document(str(path))
    sections: list[ParsedSection] = []
    heading = path.stem
    buffer: list[str] = []
    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue
        if paragraph.style and paragraph.style.name.lower().startswith("heading"):
            if buffer:
                sections.append(ParsedSection(text="\n".join(buffer), section=heading))
                buffer = []
            heading = text
        else:
            buffer.append(text)
    if buffer:
        sections.append(ParsedSection(text="\n".join(buffer), section=heading))

    for table_index, table in enumerate(document.tables, start=1):
        rows = [
            " | ".join(normalize_text(cell.text) for cell in row.cells)
            for row in table.rows
        ]
        text = "\n".join(row for row in rows if row.strip(" |"))
        if text:
            sections.append(
                ParsedSection(text=text, section=f"Table {table_index}", metadata={"table": True})
            )
    return sections


def _detect_header(rows: list[list[Any]]) -> int | None:
    expected = {"kod", "code", "sku", "ürün", "urun", "product", "fiyat", "price", "isim"}
    for index, row in enumerate(rows[:25]):
        values = [normalize_text(str(value)).lower() for value in row if value not in (None, "")]
        if len(values) >= 2 and any(any(token in value for token in expected) for value in values):
            return index
    return None


def parse_xlsx(path: Path) -> list[ParsedSection]:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=True)
    sections: list[ParsedSection] = []
    try:
        for sheet in workbook.worksheets:
            rows = [list(row) for row in sheet.iter_rows(values_only=True)]
            while rows and not any(value not in (None, "") for value in rows[-1]):
                rows.pop()
            if not rows:
                continue
            header_index = _detect_header(rows)
            if header_index is not None:
                raw_headers = rows[header_index]
                headers = [
                    normalize_text(str(value)) if value not in (None, "") else f"Column {i + 1}"
                    for i, value in enumerate(raw_headers)
                ]
                for excel_row, row in enumerate(rows[header_index + 1 :], start=header_index + 2):
                    fields = []
                    for header, value in zip(headers, row, strict=False):
                        if value not in (None, ""):
                            fields.append(f"{header}: {normalize_text(str(value))}")
                    if fields:
                        sections.append(
                            ParsedSection(
                                text="\n".join(fields),
                                section=f"{sheet.title} row {excel_row}",
                                metadata={"sheet": sheet.title, "row": excel_row},
                            )
                        )
            else:
                for start in range(0, len(rows), 25):
                    text_rows = []
                    for row in rows[start : start + 25]:
                        values = [
                            normalize_text(str(value))
                            for value in row
                            if value not in (None, "")
                        ]
                        if values:
                            text_rows.append(" | ".join(values))
                    if text_rows:
                        sections.append(
                            ParsedSection(
                                text="\n".join(text_rows),
                                section=f"{sheet.title} rows {start + 1}-{start + len(text_rows)}",
                                metadata={"sheet": sheet.title},
                            )
                        )
    finally:
        workbook.close()
    return sections


def parse_csv(path: Path) -> list[ParsedSection]:
    raw = path.read_text(encoding="utf-8-sig", errors="replace")
    try:
        dialect = csv.Sniffer().sniff(raw[:5000])
    except csv.Error:
        dialect = csv.excel
    reader = csv.DictReader(io.StringIO(raw), dialect=dialect)
    sections = []
    for row_number, row in enumerate(reader, start=2):
        fields = [
            f"{normalize_text(str(key))}: {normalize_text(str(value))}"
            for key, value in row.items()
            if key and value not in (None, "")
        ]
        if fields:
            sections.append(
                ParsedSection(
                    text="\n".join(fields),
                    section=f"Row {row_number}",
                    metadata={"row": row_number},
                )
            )
    return sections


def parse_html_text(text: str, *, section: str | None = None) -> ParsedSection:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(text, "html.parser")
    for node in soup(["script", "style", "noscript", "svg"]):
        node.decompose()
    title = normalize_text(soup.title.get_text(" ", strip=True)) if soup.title else section
    body = normalize_text(soup.get_text("\n", strip=True))
    return ParsedSection(text=body, section=title or section)


def parse_json(path: Path) -> list[ParsedSection]:
    value = json.loads(path.read_text(encoding="utf-8-sig"))
    text = json.dumps(value, ensure_ascii=False, indent=2)
    return [ParsedSection(text=text, section=path.stem)]


def parse_file(path: Path) -> list[ParsedSection]:
    extension = path.suffix.lower()
    if extension == ".pdf":
        return parse_pdf(path)
    if extension == ".docx":
        return parse_docx(path)
    if extension == ".xlsx":
        return parse_xlsx(path)
    if extension == ".csv":
        return parse_csv(path)
    if extension in {".html", ".htm"}:
        return [parse_html_text(path.read_text(encoding="utf-8-sig", errors="replace"))]
    if extension == ".json":
        return parse_json(path)
    if extension in {".txt", ".md"}:
        return [
            ParsedSection(
                text=normalize_text(path.read_text(encoding="utf-8-sig", errors="replace")),
                section=path.stem,
            )
        ]
    raise ValueError(f"Unsupported file type: {extension}")
